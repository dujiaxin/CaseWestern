# -*- coding: utf-8 -*-
"""extraction.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1Y5-2mAOWCmSegwi-jNqovJ1j07E2S6YG
"""

import re
import sys
import os
import getopt
'''Third party package'''
import pandas as pd
import docx2txt as d2t
from docx import Document
from natsort import natsorted
'''Debug used'''
import pdb, traceback, sys

'''generate new tracker file to check if every file is processed'''
process_tracked = pd.DataFrame()
'''matter id array'''
matter_id_array = []
'''rms array'''
rms_array = []
'''fot marking the process status on each file'''
check = []

text_filepath = 'C:/Users/Matt/Documents/Data Science/CW/BATCH_3/'
root_cleaned_filepath = 'C:/Users/Matt/Documents/Data Science/CW/CLEANED_3/'
if not os.path.exists(root_cleaned_filepath):
    os.makedirs(root_cleaned_filepath)

cn_reg = "Ca[a-zA-Z]*\sN[a-zA-Z]*[:;\s]"
cd_reg = "Ca[a-zA-Z]*\sD[a-zA-Z]*[:;\s]"

black_list = [
            "1st Police Report (Police Reports) 518-5137 99-86851 two versions.two victims.docx",
            "1st Police Report (Police Reports) 519-7724 06-264391.docx",
            "1st Police Report (Police Reports) 519-8861 07-242007.docx",
            "Incident Report (Public Record) (Police Reports) 512-3732 98-92533.docx",
            "Incident Report (Public Record) (Police Reports) 515-4908 00-449479.docx",
            "Incident Report (Public Record) (Police Reports) 518-3620 03-396772.docx",
            "Incident Report (Public Record) (Police Reports) 521-7536 10-176685.docx",
            "Investigative Report (Police Reports) 518-6520 06-52157.docx",
            
            
            "M5138572_RMS99-00052748.docx",
            "m5133649 rms97-5282.docx",
            "m5135901 rms94-46593 corrected to 94-46953.docx",
            "M5159011_RMS99-9226.docx",
            "M5151378_RMS07-00387676 no scu followup.docx",
            "M5124445_RMS_95-00077200_Victim 2.docx",
            "m5127211 rms96-1064701.docx",
            "M5124005_RMS 96-66727.docx",
            "m5127648 rms94-17353.docx",
            "M5141658_RMS00-00425165.docx",
            "m5144728 rms99-19755.docx",
            "M5141658_RMS00-00425165.docx",
            "m5149852 rms08-267167.docx",
            "M5178215_RMS95-44057_SUPP.REPORT.docx",
            "M5176071_RMS00-164061.docx"
        ]
# renaming file that is causing errors
# M5192379A_RMS97-100835.docx -> M5192379_RMS97-100835.docx

"""This part is a sub section extracted from next section, it is for convenience to check with the filename format is followed by rule 'Mxxx-xxxx_RMSxxx-xxxx"""
for root, dirs, files in os.walk(text_filepath,topdown=True):
        '''to skip first loop, since first loop check the current dir but we want to dive into each sub-folder'''
        '''please make sure there is no any file in first 'text_filepath' '''
        if not files:
            continue

        '''GET folder name'''
        hierachy = root.split("/")
        folder = int(hierachy[-1])
        #print(folder)

        for filename in files:

            '''extract matter_id and rms from filename'''
            filename_without_suffix = filename.split(".")[0]
            
            '''one type of file is separate by space '\s' '''
            #filename_array = filename_without_suffix.split(" ")
            filename_without_suffix = filename_without_suffix.replace(" ","")
            
            #'''if '\s' does not work, use '_' '''
            #if len(filename_array) == 1:
            filename_array = filename_without_suffix.split("_")
            
            
            matter_id_and_rms_array = []

            '''try to find matter id and rms via iterating filename array'''
            for information in filename_array:
                
                information = information.lower()
                information = information.replace(' ', '')

                if 'rms' in information:
                    information = information.replace("rms","")

                    if information.replace("-","").isdigit():
                        matter_id_and_rms_array.append(information)
                
                if 'm' in information:
                    information = information.replace("m","")
                    if information.find('-') != 4:
                        information = information[0:3] + '-' + information[3:]
                    
                    if information.replace("-","").isdigit():
                        matter_id_and_rms_array.append(information)
                    
            '''does not read files from black list'''
            if filename in black_list:
                continue

            '''the array should be a 2-length array'''
            try:
                assert len(matter_id_and_rms_array) == 2, str(folder)+": "+filename +": "+ str(filename_array)
            except AssertionError as e:
                print('folder: ' + str(folder) + '; file: ' + filename)
                exit(1)

try:
    '''start system walk to process files'''
    for root, dirs, files in os.walk(text_filepath,topdown=True):
        '''to skip first loop, since first loop check the current dir but we want to dive into each sub-folder'''
        '''please make sure there is no any file in first 'text_filepath' '''
        if not files:
            continue

        '''GET folder name'''
        hierachy = root.split("/")
        folder = int(hierachy[-1])
        #print(folder)

        cleaned_fielpath = root_cleaned_filepath+str(folder)+"/"
        if not os.path.exists(cleaned_fielpath):
            os.makedirs(cleaned_fielpath)

        for filename in files:

            '''extract matter_id and rms from filename'''
            filename_without_suffix = filename.split(".")[0]
            
            '''one type of file is separate by space '\s' '''
            #filename_array = filename_without_suffix.split(" ")
            filename_without_suffix = filename_without_suffix.replace(" ","")
            
            #'''if '\s' does not work, use '_' '''
            #if len(filename_array) == 1:
            filename_array = filename_without_suffix.split("_")
            
            
            matter_id_and_rms_array = []

            '''try to find matter id and rms via iterating filename array'''
            for information in filename_array:
                
                information = information.lower()
                information = information.replace(' ', '')

                if 'rms' in information:
                    information = information.replace("rms","")

                    if information.replace("-","").isdigit():
                        matter_id_and_rms_array.append(information)
                
                if 'm' in information:
                    information = information.replace("m","")
                    if information.find('-') != 4:
                        information = information[0:3] + '-' + information[3:]
                    
                    if information.replace("-","").isdigit():
                        matter_id_and_rms_array.append(information)
                    
            '''does not read files from black list'''
            if filename in black_list:
                #matter_id_array.append(matter_id)
                #rms_array.append(rms)
                #check.append(0)
                continue

            '''the array should be a 2-length array'''
            assert len(matter_id_and_rms_array) == 2, str(folder)+": "+filename +": "+ str(filename_array)

            '''matter_id'''
            matter_id = matter_id_and_rms_array[0]
            '''rms'''
            rms = matter_id_and_rms_array[1]
            
            

            '''filepath for each file'''
            file = root +"/"+ filename

            #print(file)
            '''read file content'''
            text_dataframe = d2t.process(file)
            

            '''search start point'''

            '''condition1 '''
            start_pointers = [sentinel.start() for sentinel in re.finditer("ORIGINAL[\D]NARRATIVE",text_dataframe, re.IGNORECASE)]

            '''condition2 '''
            if not start_pointers:
                start_pointers = [sentinel.start() for sentinel in re.finditer("SUPPLEMENTAL",text_dataframe, re.IGNORECASE)]
            
            '''condition3 '''
            if not start_pointers:
                start_pointers = [sentinel.start() for sentinel in re.finditer("facts[\D]of[\D]arrest",text_dataframe, re.IGNORECASE)]

            '''condition4 '''
            if not start_pointers:
                start_pointers = [sentinel.start() for sentinel in re.finditer("additional[\D]information",text_dataframe, re.IGNORECASE)]

            '''check again'''
            '''the array should be a 1-length array'''
            #assert len(start_pointers) !=0 , str(folder)+": "+filename +": "+ str(filename_array)+": "+str(start_pointers)
            if len(start_pointers) ==0:
                matter_id_array.append(matter_id)
                rms_array.append(rms)
                check.append(0)
                print(str(folder)+": "+filename +": "+ str(filename_array)+": "+str(start_pointers)+" NO TITLE CONDITION MEET!")
                continue
            elif len(start_pointers) > 1:
                matter_id_array.append(matter_id)
                rms_array.append(rms)
                check.append(0)
                print(str(folder)+": "+filename +": "+ str(filename_array)+": "+str(start_pointers)+" "+str(len(start_pointers))+" TITLE CONDITION MEET!")
                continue


            '''where text starts'''
            start_pointer = start_pointers[-1]
            #del text_dataframe
            
            '''start to process data'''

            '''position starts from "ORIGINAL NARRATIVE" '''
            head = start_pointer

            #print("map index is "+str(map_index))
            #print(mater_id)

            '''first line for every processed document'''
            heading_string = 'Mater ID: '+str(matter_id)+" RMS: "+str(rms)+"\n"

            ''' "text" contains text between "head" and "tail" from each docs '''
            text = text_dataframe[head:]
            text_array = text.split("\n")
            pure_text_array = [i.lstrip().rstrip() for i in text_array if i]
            #print(pure_text_array)

            processed_text_array = []
            extra_followed = False
            check_extra_followed_array = []

            #print(pure_text_array)
            '''to filtering data'''
            for t in pure_text_array:

                if extra_followed == True:
                    extra_followed = False
                    continue

                if re.search(cd_reg,t,re.IGNORECASE):
                    check_extra_followed_array = t.split(" ")
                    check_extra_followed_array = [i for i in check_extra_followed_array if i]
                    if len(check_extra_followed_array)!=2:
                        extra_followed = True
                    continue

                if re.search(cn_reg,t,re.IGNORECASE):
                    check_extra_followed_array = t.split(" ")
                    check_extra_followed_array = [i for i in check_extra_followed_array if i]
                    if len(check_extra_followed_array)!=2:
                        extra_followed = True
                    continue

                if "OFFENSE/INCIDENT".lower() in t.lower():
                    #import ipdb; ipdb.set_trace() # debugging starts here
                    continue

                if "CLEVELAND POLICE DEPARTMENT".lower() in t.lower():
                    continue     

                if "Date:".lower() in t.lower():
                    check_extra_followed_array = t.split(":")
                    if len(check_extra_followed_array)!=2:
                        extra_followed = True
                    continue    

                if "Page:".lower() in t.lower():
                    check_extra_followed_array = t.split(":")
                    if len(check_extra_followed_array)!=2:
                        extra_followed = True
                    continue

                if "From: unknown".lower() in t.lower():
                    continue
                if "This fax".lower() in t.lower():
                    continue

                processed_text_array.append(t)

            document = Document()
            document.add_paragraph(heading_string)
            for line in processed_text_array:
                document.add_paragraph(line)
            document.save(cleaned_fielpath+"M"+str(matter_id)+"_"+"RMS"+str(rms)+".docx")

            '''tracker'''
            matter_id_array.append(matter_id)
            rms_array.append(rms)
            check.append(1)
            
            del pure_text_array
            del document
            del text
            del text_dataframe
                

        #print(len(mater_id_array))
        #print(len(rms_array))
        #print(len(check))
    process_tracked['Matter_id']  = matter_id_array
    process_tracked['RMS'] = rms_array
    process_tracked['processed'] = check

    process_tracked.to_csv(root_cleaned_filepath+"document_process.csv")
    print("Finished!")
except:
        extype, value, tb = sys.exc_info()
        traceback.print_exc()
        pdb.post_mortem(tb)